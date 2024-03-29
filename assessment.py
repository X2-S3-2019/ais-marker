# -*- coding: utf-8 -*-
# Developer's Guide
# For every Assessment Document, a template is used.
# For every template, there are criteria. Criteria may be grouped.
# For every criterion, there are fields.


import json
import os
import sys
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm


class Field:
    def __init__(self, id, name, value, description, points):
        self.id = id
        self.name = name
        self.value = value
        self.description = description
        self.points = points


class Criterion:
    def __init__(self, id, name, description,  icon):
        self.id = id
        self.name = name
        self.description = description
        self.icon = icon
        self.fields = []

    def addField(self, field):
        self.fields.append(field)

    def getFields(self):
        return self.fields

    def printFields(self):
        for i in self.fields:
            print(i.description)


class GroupCriterion:
    def __init__(self, id, name, icon):
        self.id = id
        self.name = name
        self.icon = icon
        self.criteria = []

    def addCriterion(self, criterion):
        self.criteria.append(criterion)

    def printCriteria(self):
        for i in self.criteria:
            print(i.name + " from Group Criterion def")

    def getCriteria(self):
        return self.criteria

    def getFirstCriterion(self):
        return self.criteria[0]


class Template:
    def __init__(self, id, name, description):
        self.id = id
        self.name = name
        self.description = description
        # If groupCriteria is null or empty, then the template only contains criteria without grouping
        self.groupCriteria = []

    def addCriterion(self, criterion):
        self.criteria.append(criterion)

    def addGroupCriterion(self, groupCriterion):
        self.groupCriteria.append(groupCriterion)

    def addCriterionToGroup(self, criterion, groupCriterion):
        # If the group criterion exists in the list, add the criterion to group
        if groupCriterion in self.groupCriteria:
            i = self.groupCriteria.index(groupCriterion)
            self.groupCriteria[i].addCriterion(criterion)

    def getGroupCriteria(self):
        return self.groupCriteria

    def getCriteria(self):
        return self.criteria

    def printCriteria(self):
        for i in self.criteria:
            print(i.name)

    def printGroupCriteria(self):
        for i in self.groupCriteria:
            print(i.name)


class AssessmentDocument:
    def __init__(self, document_name):
        self.document_name = document_name

    def openResultsDocument(self):
        os.startfile(self.document_name + '.docx')

    # Create document containing assessment results using a template
    # Add shading of selected marking

    def createGroupResultsDocument(self, template, headerInfo, groupData, scoreData, path):
        document = Document()
        print('Path: ' + path)

        # Change the page margins
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.75)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(2)

        # Add the date of the presentation
        date_header = document.add_paragraph()
        date_header.add_run(
            'Presentation date: ' + headerInfo['presentationDate'])
        date_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Iterate through student names and marks
        for student in groupData['members']:
            user_id = str(student['id'])
            
            studentScore = scoreData[user_id]['assessmentTotalScore']

            print('---scoreData[id]---')
            print(scoreData[user_id])
            
            table = document.add_table(1, 3)
            assessment_info_cells = table.rows[0].cells
            student_name_cell = assessment_info_cells[0]
            student_name_cell.paragraphs[0].add_run("Student " + user_id + " name: ")
            student_name_cell.paragraphs[0].add_run(
                student['studentName']).bold = True
            student_id_cell = assessment_info_cells[1]
            student_id_cell.paragraphs[0].add_run("Student ID: ")
            student_id_cell.paragraphs[0].add_run(
                student['studentId']).bold = True
            student_marks_cell = assessment_info_cells[2]
            student_marks_cell.paragraphs[0].add_run("Obtained Marks: ")
            student_marks_cell.paragraphs[0].add_run(str(studentScore) + "/" + str(scoreData['perfectScore'])).bold = True

        # Add a new line
        document.add_paragraph()

        # For each group criterion, create a table
        groupCriteria = template.getGroupCriteria()
        for group in groupCriteria:
            table = document.add_table(1, 5, style='Table Grid')
            # Get the first row of the table, which contains the headers
            heading_cells = table.rows[0].cells
            # Add the group criterion's name
            heading_cells[0].paragraphs[0].add_run(group.name).bold = True
            # Add value header
            first_criterion_fields = group.getFirstCriterion().getFields()
            # group.getFirstCriterion().printFields()
            ctr = 1
            # For each field in the first criterion, add the value to header
            for field in first_criterion_fields:
                heading_cells[ctr].paragraphs[0].add_run(
                    str(field.points) + ' - ' + field.value).bold = True
                ctr += 1

            # Used to access the results object
            temp_words = group.name.split(' ')
            data_type_group = temp_words[0].lower() + '_' + str(group.id)

            criteria = group.getCriteria()
            for criterion in criteria:
                criterion_fields = criterion.getFields()

                ctr = 0
                cells = table.add_row().cells
                cells[ctr].text = criterion.name

                # Used to access the results object
                temp_words = criterion.name.split(' ')
                data_type_criterion = temp_words[0].lower(
                ) + '_' + str(criterion.id)

                for field in criterion_fields:
                    ctr += 1
                    #cells[ctr].text = field.description

                    cell_text = field.description
                   

                    # Iterate through students' results
                    for student in groupData['members']:
                        user_id = str(student['id'])
                        score = scoreData[user_id]['groupCriteria'][data_type_group]['criteria'][data_type_criterion]
                        if(str(score) == str(field.points)):
                            # Add the student index
                            cell_text = cell_text + '(' + str(student['id']) +')'
                    
                    cells[ctr].text = cell_text

            # Add a new line for every group criterion
            document.add_paragraph()

        footer = document.add_paragraph().add_run(
            'Adapted with enhancement from © 2004 National Council of Teachers of English/International Reading Association')
        font = footer.font
        font.size = Pt(8)
        font.italic = True

        document.save(path + '/' + self.document_name + '.docx')


    def createResultsDocument(self, template, results, otherInfo, path):
        document = Document()
        print('Path: ' + path)

        # Change the page margins to match Saghir's SOFT808
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.75)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(2)

        studentScore = ''
        perfectScore = ''
        scorePercentage = str(results['scorePercentage'])
        if results['normalizedScore'] != '':
            studentScore = str(results['normalizedScore'])
            perfectScore = str(results['normalizedPerfectScore'])
        else:
            studentScore = str(results['assessmentTotalScore'])
            perfectScore = str(results['assessmentPossibleTotalScore'])

        date_header = document.add_paragraph()
        date_header.add_run(
            'Presentation date: ' + otherInfo['presentationDate'])
        date_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add presentation duration
        print(otherInfo['presentationDuration'])
        if otherInfo['presentationDuration'] != '':
            duration_header = document.add_paragraph()
            duration_header.add_run(
                'Presentation duration: ' + otherInfo['presentationDuration'])
            duration_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # TODO: Make this dynamic/customizable for each template
        # Add the student name, id and topic of presentation

        table = document.add_table(1, 3)

        assessment_info_cells = table.rows[0].cells
        student_name_cell = assessment_info_cells[0]
        student_name_cell.paragraphs[0].add_run("Student name: ")
        student_name_cell.paragraphs[0].add_run(
            otherInfo['studentName']).bold = True
        student_id_cell = assessment_info_cells[1]
        student_id_cell.paragraphs[0].add_run("Student ID: ")
        student_id_cell.paragraphs[0].add_run(
            otherInfo['studentId']).bold = True
        student_marks_cell = assessment_info_cells[2]
        student_marks_cell.paragraphs[0].add_run("Obtained Marks: ")
        #student_marks_cell.paragraphs[0].add_run(studentScore + "/" + perfectScore + " (" + scorePercentage + "%)").bold = True
        student_marks_cell.paragraphs[0].add_run(studentScore + "/" + perfectScore).bold = True

        table = document.add_table(1, 1)
        assessment_info_cells = table.rows[0].cells
        topic_of_presentation_cell = assessment_info_cells[0]
        topic_of_presentation_cell.paragraphs[0].add_run(
            "Topic of Presentation: ")
        topic_of_presentation_cell.paragraphs[0].add_run(
            otherInfo['presentationTopic']).bold = True

        # Add a new line
        document.add_paragraph()

        # For each group criterion, create a table
        groupCriteria = template.getGroupCriteria()
        for group in groupCriteria:
            table = document.add_table(1, 5, style='Table Grid')
            # Get the first row of the table, which contains the headers
            heading_cells = table.rows[0].cells
            # Add the group criterion's name
            heading_cells[0].paragraphs[0].add_run(group.name).bold = True
            # Add value header
            first_criterion_fields = group.getFirstCriterion().getFields()
            # group.getFirstCriterion().printFields()
            ctr = 1
            # For each field in the first criterion, add the value to header
            for field in first_criterion_fields:
                heading_cells[ctr].paragraphs[0].add_run(
                    str(field.points) + ' - ' + field.value).bold = True
                ctr += 1

            # Used to access the results object
            temp_words = group.name.split(' ')
            data_type_group = temp_words[0].lower() + '_' + str(group.id)

            criteria = group.getCriteria()
            for criterion in criteria:
                criterion_fields = criterion.getFields()

                ctr = 0
                cells = table.add_row().cells
                cells[ctr].text = criterion.name

                # Used to access the results object
                temp_words = criterion.name.split(' ')
                data_type_criterion = temp_words[0].lower(
                ) + '_' + str(criterion.id)

                for field in criterion_fields:
                    ctr += 1
                    cells[ctr].text = field.description

                    # Check if the field was selected based on the results
                    # TODO: Make a colored version of shading
                    score = results['groupCriteria'][data_type_group]['criteria'][data_type_criterion]
                    if (score == str(field.points)):
                        # Set a cell background (shading) color to RGB D9D9D9.
                        shading_elm = parse_xml(
                            r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                        cells[ctr]._tc.get_or_add_tcPr().append(shading_elm)

            document.add_paragraph()

            groupTotal = document.add_paragraph('(' + group.name + ') ' + 'Total: ' + str(
                results['groupCriteria'][data_type_group]['groupTotalScore']) + ' / ' + str(results['groupCriteria'][data_type_group]['groupPossibleTotalScore']))
            groupTotal.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add a new line for every group criterion
            document.add_paragraph()

        footer = document.add_paragraph().add_run(
            'Adapted with enhancement from © 2004 National Council of Teachers of English/International Reading Association')
        font = footer.font
        font.size = Pt(8)
        font.italic = True

        document.save(path + '/' + self.document_name + '.docx')


class TemplateDocument:
    def __init__(self, name):
        self.name = name

    def createTemplateDocument(self, template):
        # Create a document with table
        document = Document()

        # For each group criterion, create a table
        groupCriteria = template.getGroupCriteria()
        for group in groupCriteria:
            table = document.add_table(1, 5, style='Table Grid')
            # Get the first row of the table, which contains the headers
            heading_cells = table.rows[0].cells
            # Add the group criterion's name
            heading_cells[0].paragraphs[0].add_run(group.name).bold = True
            # Add value header
            first_criterion_fields = group.getFirstCriterion().getFields()
            # group.getFirstCriterion().printFields()
            ctr = 1
            # For each field in the first criterion, add the value to header
            for field in first_criterion_fields:
                heading_cells[ctr].paragraphs[0].add_run(
                    str(field.points) + ' - ' + field.value).bold = True
                ctr += 1

            criteria = group.getCriteria()
            for criterion in criteria:
                criterion_fields = criterion.getFields()

                ctr = 0
                cells = table.add_row().cells
                cells[ctr].text = criterion.name
                for field in criterion_fields:
                    ctr += 1
                    cells[ctr].text = field.description

            # Add a new line for every group criterion
            document.add_paragraph()

        document.save(self.name + '.docx')

    def openTemplateDocument(self):
        os.startfile(self.name + '.docx')
